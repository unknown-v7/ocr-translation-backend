import os
import requests
from flask import Flask, request, jsonify, send_file
import base64
import io
from io import BytesIO
from PIL import Image
import pytesseract
from deep_translator import GoogleTranslator
import datetime
import sqlite3
from gtts import gTTS
from reportlab.pdfgen import canvas
import cv2
import numpy as np
from reportlab.lib.pagesizes import letter
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
import docx.shared
from openai import OpenAI
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
app = Flask(__name__)

def enhance_image(pil_img):
    img = np.array(pil_img.convert('L'))  # 转灰度
    img = cv2.adaptiveThreshold(img, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                                cv2.THRESH_BINARY, 15, 10)
    return Image.fromarray(img)

# --- 初始化数据库 ---
def init_db():
    conn = sqlite3.connect('history.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT NOT NULL,
            ocr_text TEXT NOT NULL,
            translation TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

init_db()

# --- 保存历史记录 ---
def save_history(ocr_text, translation):
    conn = sqlite3.connect('history.db')
    c = conn.cursor()
    c.execute('INSERT INTO history (timestamp, ocr_text, translation) VALUES (?, ?, ?)',
              (datetime.datetime.now().isoformat(), ocr_text, translation))
    conn.commit()
    conn.close()

# --- 获取历史记录 ---
def get_history(limit=50):
    conn = sqlite3.connect('history.db')
    c = conn.cursor()
    c.execute('SELECT id, timestamp, ocr_text, translation FROM history ORDER BY id DESC LIMIT ?', (limit,))
    rows = c.fetchall()
    conn.close()
    return [{'id': row[0], 'timestamp': row[1], 'ocr_text': row[2], 'translation': row[3]} for row in rows]

# --- 获取历史统计数据 ---
def get_history_stats():
    conn = sqlite3.connect('history.db')
    c = conn.cursor()

    # 总记录数
    c.execute('SELECT COUNT(*) FROM history')
    total_records = c.fetchone()[0]

    # 最近7天使用趋势
    c.execute('''
        SELECT DATE(timestamp) as date, COUNT(*) as count
        FROM history
        WHERE timestamp >= date('now', '-7 days')
        GROUP BY date
        ORDER BY date
    ''')
    usage_trend = [{'date': row[0], 'count': row[1]} for row in c.fetchall()]

    # 计算平均词数（简单用空格数+1估算）
    c.execute('''
        SELECT (LENGTH(ocr_text) - LENGTH(REPLACE(ocr_text, ' ', '')) + 1) AS word_count,
               (LENGTH(translation) - LENGTH(REPLACE(translation, ' ', '')) + 1) AS trans_word_count
        FROM history
    ''')
    word_counts = c.fetchall()
    avg_words = sum(row[0] for row in word_counts) / len(word_counts) if word_counts else 0
    avg_trans_words = sum(row[1] for row in word_counts) / len(word_counts) if word_counts else 0

    conn.close()

    return {
        'total_records': total_records,
        'usage_trend': usage_trend,
        'avg_words': round(avg_words, 1),
        'avg_trans_words': round(avg_trans_words, 1)
    }

# --- base64转图片 ---
def base64_to_image(base64_str):
    img_data = base64.b64decode(base64_str)
    return Image.open(BytesIO(img_data))


@app.route('/', methods=['GET'])
def hello():
    return "OCR 翻译系统后端运行正常"


@app.route('/ocr', methods=['POST'])
def ocr_and_translate():
    try:
        data = request.json
        images_base64 = data.get("images_base64")
        ocr_lang = data.get("ocr_lang", "eng")
        target_lang = data.get("target_lang", "zh-CN")

        if not images_base64:
            return jsonify({"error": "图像数据为空"}), 400
        if not isinstance(images_base64, list):
            images_base64 = [images_base64]

        results = []
        for img_b64 in images_base64:
            image = base64_to_image(img_b64)
            text = pytesseract.image_to_string(image, lang=ocr_lang).strip()
            if not text:
                text = "[未识别出文字]"
            translation = GoogleTranslator(source='auto', target=target_lang).translate(text)
            save_history(text, translation)

            results.append({
                "ocr_text": text,
                "translation": translation
            })

        return jsonify({"results": results})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/history', methods=['GET'])
def history():
    limit = request.args.get('limit', 50, type=int)
    records = get_history(limit)
    return jsonify({"history": records})


@app.route('/stats', methods=['GET'])
def stats():
    stats_data = get_history_stats()
    return jsonify(stats_data)

# 从环境变量读取 API 密钥
ARK_API_KEY = os.getenv("ARK_API_KEY")
if not ARK_API_KEY:
    raise Exception("请设置环境变量 ARK_API_KEY")

client = OpenAI(
    base_url="https://ark.cn-beijing.volces.com/api/v3",
    api_key=os.getenv("ARK_API_KEY")
)

def call_doubao_ai_summary(text, prompt="请对以下内容进行总结并优化表达："):
    response = client.chat.completions.create(
        model="ep-20250703150551-l6jrl",  # 你的推理服务 ID
        messages=[
            {"role": "system", "content": "你是一个语言精炼助手。"},
            {"role": "user", "content": f"{prompt}\n\n{text}"}
        ]
    )
    return response.choices[0].message.content

# Flask 路由：POST /ai_summary
@app.route("/ai_summary", methods=["POST"])
def ai_summary():
    try:
        data = request.json
        text = data.get("text", "")
        prompt = data.get("prompt", "请对以下内容进行总结并优化表达：")

        if not text.strip():
            return jsonify({"error": "缺少文本内容"}), 400

        summary = call_doubao_ai_summary(text, prompt)
        return jsonify({"summary": summary})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/tts', methods=['POST'])
def text_to_speech():
    data = request.json
    text = data.get('text', '')
    lang = data.get('lang', 'zh-CN')

    if not text:
        return jsonify({"error": "文本内容为空"}), 400

    try:
        tts = gTTS(text=text, lang=lang)
        audio_buffer = BytesIO()
        tts.write_to_fp(audio_buffer)
        audio_buffer.seek(0)

        return send_file(
            audio_buffer,
            mimetype='audio/mpeg',
            as_attachment=True,
            download_name='speech.mp3'
        )
    except Exception as e:
        return jsonify({"error": f"语音生成失败: {str(e)}"}), 500


@app.route('/export_pdf', methods=['GET'])
def export_pdf():
    limit = request.args.get('limit', 50, type=int)
    records = get_history(limit)

    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    p.setFont("Helvetica-Bold", 16)
    p.drawCentredString(width / 2, height - 50, "OCR翻译历史记录")

    p.setFont("Helvetica-Bold", 12)
    p.drawString(50, height - 80, "时间")
    p.drawString(200, height - 80, "识别内容")
    p.drawString(400, height - 80, "翻译内容")

    y = height - 100
    for record in records:
        if y < 100:
            p.showPage()
            y = height - 50
            p.setFont("Helvetica-Bold", 12)
            p.drawString(50, y, "时间")
            p.drawString(200, y, "识别内容")
            p.drawString(400, y, "翻译内容")
            y -= 30

        p.setFont("Helvetica", 10)
        p.drawString(50, y, record['timestamp'][:19])

        # 长文本换行处理
        def split_text(text, length):
            return [text[i:i+length] for i in range(0, len(text), length)]

        ocr_lines = split_text(record['ocr_text'], 60)
        trans_lines = split_text(record['translation'], 30)
        max_lines = max(len(ocr_lines), len(trans_lines))

        for i in range(max_lines):
            if i > 0:
                y -= 15
                if y < 50:
                    p.showPage()
                    y = height - 50

            if i < len(ocr_lines):
                p.drawString(200, y, ocr_lines[i])
            if i < len(trans_lines):
                p.drawString(400, y, trans_lines[i])

        y -= 30

    p.save()
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="ocr_history.pdf", mimetype='application/pdf')


@app.route('/export_word', methods=['GET'])
def export_word():
    limit = request.args.get('limit', 50, type=int)
    records = get_history(limit)

    doc = Document()
    doc.add_heading('OCR翻译历史记录', level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '时间'
    hdr_cells[1].text = '识别内容'
    hdr_cells[2].text = '翻译内容'

    for record in records:
        row_cells = table.add_row().cells
        row_cells[0].text = record['timestamp'][:19]
        row_cells[1].text = record['ocr_text']
        row_cells[2].text = record['translation']

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name='ocr_history.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/export_excel', methods=['GET'])
def export_excel():
    limit = request.args.get('limit', 50, type=int)
    records = get_history(limit)

    df = pd.DataFrame(records)
    df = df[['timestamp', 'ocr_text', 'translation']]
    df['timestamp'] = df['timestamp'].str[:19]

    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='翻译历史')
        workbook = writer.book
        worksheet = writer.sheets['翻译历史']
        worksheet.set_column('A:A', 20)
        worksheet.set_column('B:B', 50)
        worksheet.set_column('C:C', 50)

    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name='ocr_history.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


@app.route('/export_stats', methods=['GET'])
def export_stats():
    stats_data = get_history_stats()

    dates = [day['date'] for day in stats_data['usage_trend']]
    counts = [day['count'] for day in stats_data['usage_trend']]

    plt.figure(figsize=(10, 5))
    plt.plot(dates, counts, marker='o')
    plt.title('最近7天使用趋势')
    plt.xlabel('日期')
    plt.ylabel('使用次数')
    plt.xticks(rotation=45)
    plt.tight_layout()

    chart_buffer = BytesIO()
    plt.savefig(chart_buffer, format='png')
    plt.close()
    chart_buffer.seek(0)

    doc = Document()
    doc.add_heading('OCR翻译系统统计报告', level=1)

    doc.add_heading('系统概览', level=2)
    doc.add_paragraph(f"总翻译记录数: {stats_data['total_records']}")
    doc.add_paragraph(f"平均原文长度: {stats_data['avg_words']} 词")
    doc.add_paragraph(f"平均译文长度: {stats_data['avg_trans_words']} 词")

    doc.add_heading('最近7天使用趋势', level=2)
    doc.add_picture(chart_buffer, width=docx.shared.Inches(6))

    doc.add_heading('每日使用数据', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '日期'
    hdr_cells[1].text = '使用次数'

    for day in stats_data['usage_trend']:
        row_cells = table.add_row().cells
        row_cells[0].text = day['date']
        row_cells[1].text = str(day['count'])

    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return send_file(
        doc_buffer,
        as_attachment=True,
        download_name='translation_stats.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5050, debug=True)


